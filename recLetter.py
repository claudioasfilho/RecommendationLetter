import Listofthings

from Listofthings import PronoumsList,PositivePersonalityTraits,AcademicSkills, Phrase1, Phrase2, Phrase3, Phrase4, Phrase5, LinkingWords
import openpyxl
from openpyxl import Workbook
from openpyxl import load_workbook
import sys, getopt
import os
import random
import datetime
from datetime import date
import docx
from docx.shared import Length,Inches, Pt

doc = docx.Document('Template.docx')

all_paras = doc.paragraphs

#it reads the input file
StudentProfile_xls = "StudentProfile.xlsx"


if os.path.isfile(StudentProfile_xls):
    #it Imports the Excel File
    wb = load_workbook(StudentProfile_xls)
    sheets = wb.sheetnames
    ws= wb[sheets[0]]

    teachersName = ws['B2'].value
    firstName = ws['B3'].value
    lastName = ws['B4'].value
    pronoum = ws['B5'].value
    pronoumLocalList = PronoumsList.get(pronoum)
    #he
    subjective = pronoumLocalList[0]
    #him
    objective = pronoumLocalList[1]
    #his
    possessive = pronoumLocalList[2]

    classAttended = ws['B6'].value

    highSchoolYearAttended = ws['B8'].value


    ###########################
    #Time teacher know students
    ###########################

    schoolYearAttended = ws['B7'].value
    yearsAttended = schoolYearAttended.split('/')
    start_date = datetime.datetime(int(yearsAttended[0]), 8, 15)
    num_months = (date.today().year - start_date.year) * 12 + (date.today().month - start_date.month)

    targettedInstitution = ws['B9'].value

    if targettedInstitution == " ":
         targettedInstitution = "your"
    else :
        targettedInstitution = "the " + ws['B9'].value


    #Column A
    colA = ws['A']
    #Column B
    colB = ws['B']


    ############################
    #Collecting purposeOfTheLetter
    ############################
    dataAvailable = 0

    #Initial Row where data will be examined, The offset needs to be set at -1
    rowOffSet = 12 - 1
    #Last Row where data will be examined
    lastRow = 16
    #counter
    rowCounter = rowOffSet
    #flag that indicates Selection was made
    dataAvailable = 0

    while dataAvailable == 0 and rowCounter<lastRow:
        #print(rowCounter)
        #print(colA[rowCounter].value)
        #print(colB[rowCounter].value)
        if dataAvailable == 0 and colB[rowCounter].value == 'X':
            purposeOfTheLetter = colA[rowCounter].value
            dataAvailable = 1
            break
        rowCounter = rowCounter + 1


    ############################
    #Collecting accomplishments
    ############################
    accomplishments = 0
    #Initial Row where data will be examined, The offset needs to be set at -1
    rowOffSet = 20 - 1
    #Last Row where data will be examined
    lastRow = 23
    #counter
    rowCounter = rowOffSet
    #flag that indicates Selection was made
    dataAvailable = 0

    while dataAvailable == 0 and rowCounter<lastRow:
        #print(rowCounter)
        #print(colA[rowCounter].value)
        #print(colB[rowCounter].value)
        if dataAvailable == 0 and colB[rowCounter].value == 'X':
            accomplishments = colA[rowCounter].value
            dataAvailable = 1
            break
        rowCounter = rowCounter + 1


    #######################################
    #Collecting Positive Personality Traits
    ######################################

    positivePersonalityTraits = []
    max_number_of_traits = 6
    #Initial Row where data will be examined, The offset needs to be set at -1
    rowOffSet = 27 - 1
    #Last Row where data will be examined
    lastRow = 61
    #counter
    rowCounter = rowOffSet
    #flag that indicates Selection was made
    traitCounter = 0

    while traitCounter <= max_number_of_traits and rowCounter<lastRow:
        #print(rowCounter)
        #print(colA[rowCounter].value)
        #print(colB[rowCounter].value)
        if colB[rowCounter].value == 'X':
            positivePersonalityTraits.append(colA[rowCounter].value)
            traitCounter = traitCounter + 1
        rowCounter = rowCounter + 1

    ############################
    #Collecting Academic Skills
    ############################

    academicSkills = []
    max_number_of_skills = 6
    #Initial Row where data will be examined, The offset needs to be set at -1
    rowOffSet = 66 - 1
    #Last Row where data will be examined
    lastRow = 80
    #counter
    rowCounter = rowOffSet
    #flag that indicates Selection was made
    skillCounter = 0

    while skillCounter <= max_number_of_skills and rowCounter<lastRow:
        if colB[rowCounter].value == 'X':
            academicSkills.append(colA[rowCounter].value)
            skillCounter = skillCounter + 1
        rowCounter = rowCounter + 1

else:
    print("{0} does not appear to be a valid file, choose the right file and retry".format(inputfile))
    sys.exit(2)

#Ramdomizing Academic Skills and Positive traits


academicSkillsFinal = []

acadSkill1 = random.choice(academicSkills)
academicSkills.remove(acadSkill1)
academicSkillsFinal.append(acadSkill1)
acadSkill2 = random.choice(academicSkills)
academicSkills.remove(acadSkill2)
academicSkillsFinal.append(acadSkill2)
acadSkill3 = random.choice(academicSkills)
academicSkills.remove(acadSkill3)
academicSkillsFinal.append(acadSkill3)

if skillCounter>3:
    acadSkill4 = random.choice(academicSkills)
    academicSkills.remove(acadSkill4)
    academicSkillsFinal.append(acadSkill4)
    if skillCounter>4:
        acadSkill5 = random.choice(academicSkills)
        academicSkills.remove(acadSkill5)
        academicSkillsFinal.append(acadSkill5)
        if skillCounter>5:
            acadSkill6 = random.choice(academicSkills)
            academicSkills.remove(acadSkill6)
            academicSkillsFinal.append(acadSkill6)


positivePersonalityTraitFinal = []
personalTrait1 = random.choice(positivePersonalityTraits)
positivePersonalityTraits.remove(personalTrait1)
positivePersonalityTraitFinal.append(personalTrait1)
personalTrait2 = random.choice(positivePersonalityTraits)
positivePersonalityTraits.remove(personalTrait2)
positivePersonalityTraitFinal.append(personalTrait2)
personalTrait3 = random.choice(positivePersonalityTraits)
positivePersonalityTraits.remove(personalTrait3)
positivePersonalityTraitFinal.append(personalTrait3)

if traitCounter>3:
    personalTrait4 = random.choice(positivePersonalityTraits)
    positivePersonalityTraits.remove(personalTrait4)
    positivePersonalityTraitFinal.append(personalTrait4)
    if traitCounter>4:
        personalTrait5 = random.choice(positivePersonalityTraits)
        positivePersonalityTraits.remove(personalTrait5)
        positivePersonalityTraitFinal.append(personalTrait5)
        if traitCounter>5:
            personalTrait6 = random.choice(positivePersonalityTraits)
            positivePersonalityTraits.remove(personalTrait6)
            positivePersonalityTraitFinal.append(personalTrait6)



_space = " "

# style = doc.styles.add_style('Indent', WD_STYLE_TYPE.PARAGRAPH)
# paragraph_format = style.paragraph_format
# # paragraph_format.left_indent = Inches(0.25)
# # paragraph_format.first_line_indent = Inches(-0.25)
# paragraph_format.line_spacing_rule = Pt(10)

################
#1st Paragraph
################

doc.add_paragraph("To whom it may concern,\n")

#print(random.choice(Phrase1) + firstName + _space + lastName + " to the " + targettedInstitution + _space)

#_1st_P = str(random.choice(Phrase1) + firstName + _space + lastName + " to the " + targettedInstitution + _space + purposeOfTheLetter + ". " + random.choice(Phrase2) + _space + objective.lower() + _space + "in my classroom. "+ firstName + " was a " + highSchoolYearAttended + " in my " + classAttended + " class in " + schoolYearAttended + ". ")
                #"It is with pleasure that I recommend "                        "I was fortunate to have" "him"
_1st_paragrapth = doc.add_paragraph (random.choice(Phrase1) + firstName + _space + lastName + " to " + targettedInstitution + _space + purposeOfTheLetter + ". " + random.choice(Phrase2) + _space + objective.lower() + _space + "in my classroom. "+ firstName + " was a " + highSchoolYearAttended + " in my " + classAttended + " class in " + schoolYearAttended + ". ")
#_1st_paragrapth = doc.add_paragraph (_1st_P)
if(num_months < 12) : _1st_paragrapth.add_run( "Although I have only taught " + firstName + _space + "for " + str(int(num_months)) + " months, I can already see ")
if(num_months > 12 and num_months < 24) : _1st_paragrapth.add_run("I have known " + firstName + _space + "for over an year, and " + subjective.lower() + " made an impression in me due to ")
if(num_months > 24) : _1st_paragrapth.add_run("I have known " + firstName + _space + "for more than " + str(int(num_months/12)) + " years, and " + subjective.lower() + " made an impression in me due to ")

#print(possessive.lower() + _space + positivePersonalityTraits[0] + _space + "and also " + positivePersonalityTraits[1] + " personality.", end="", flush=True)
_1st_paragrapth.add_run(objective.lower() + _space + random.choice(academicSkillsFinal) + _space + "and also " +  random.choice(positivePersonalityTraitFinal) + " personality. ",)

#"I want to illustrate a little more about ",         #Him/Her
_1st_paragrapth.add_run(random.choice(Phrase3) +  objective.lower() + " in this letter, and why " + subjective.lower() +  " deserves to be considered in your instituition.")


doc.add_paragraph("\n")

################
#2nd Paragraph
################

acadSkillA = random.choice(academicSkillsFinal)
academicSkillsFinal.remove(acadSkillA)

acadSkillB = random.choice(academicSkillsFinal)
academicSkillsFinal.remove(acadSkillB)

acadSkillC = random.choice(academicSkillsFinal)
academicSkillsFinal.remove(acadSkillC)

# acadSkillD = random.choice(academicSkillsFinal)
# academicSkillsFinal.remove(acadSkillD)


doc.add_paragraph("While in class I have observed some remarkable academic skills. " + firstName + _space + AcademicSkills[acadSkillA] + ". " + subjective  + " also " + AcademicSkills[acadSkillB]
+ random.choice(Phrase5) + subjective.lower() + _space + AcademicSkills[acadSkillC] + ".")#+ random.choice(LinkingWords) + subjective.lower() + _space  + AcademicSkills[acadSkillD] + ".")

doc.add_paragraph("\n")


################
#3rd Paragraph
################


personalTraitA = random.choice(positivePersonalityTraitFinal)
positivePersonalityTraitFinal.remove(personalTraitA)

personalTraitB = random.choice(positivePersonalityTraitFinal)
positivePersonalityTraitFinal.remove(personalTraitB)

personalTraitC = random.choice(positivePersonalityTraitFinal)
positivePersonalityTraitFinal.remove(personalTraitC)

# personalTraitD = random.choice(positivePersonalityTraitFinal)
# positivePersonalityTraitFinal.remove(personalTraitD)


doc.add_paragraph("Besides all " + possessive.lower() + " Academic work, " + firstName + _space + "is a very "+ personalTraitA + " student. " + subjective + _space +  PositivePersonalityTraits[personalTraitB] + ". " + subjective  + " also " + PositivePersonalityTraits[personalTrait2]
+ random.choice(Phrase5) + subjective.lower() + _space + PositivePersonalityTraits[personalTraitC]+ ".")#+ random.choice(LinkingWords) + subjective.lower() + _space  + PositivePersonalityTraits[personalTraitD] + ".")

doc.add_paragraph("\n")

lastParagraph = doc.add_paragraph("Please, contact me if you have any questions.\n")
lastParagraph.add_run("Sincerely,\n\n")

lastParagraph.add_run(teachersName + "\n")
lastParagraph.add_run(str(date.today().month) + "/" + str(date.today().day) + "/" + str(date.today().year))

file_name = (firstName + lastName + "-" + str(date.today().month) + "-" + str(date.today().day) + "-" + str(date.today().year) + ".docx")
doc.save(file_name)

print(file_name + " created sucessfully")
